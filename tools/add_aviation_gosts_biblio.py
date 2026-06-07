from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/kotewa/my/diplom на проверку.docx")

GOST_ITEMS = [
    "ГОСТ Р 55585-2013. Воздушный транспорт. Система управления безопасностью полетов воздушных судов. Термины и определения [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/9e224af2-9c65-4ccc-9d35-fa2dcc116ef5 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 55588-2013. Воздушный транспорт. Система менеджмента безопасности авиационной деятельности. Термины и определения [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/e53c8776-82a7-4f73-ac0f-3545cfa10d14 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 56079-2014. Изделия авиационной техники. Безопасность полета, надежность, контролепригодность, эксплуатационная и ремонтная технологичность. Номенклатура показателей [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/fc9bc901-8f2c-40f5-9ada-b100881a43c5 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 57240-2016. Воздушный транспорт. Менеджмент безопасности авиационной деятельности в гражданской авиации. Основные положения [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/6094796d-8ed9-48b0-b222-09e48ffc05f0 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 56494-2015. Воздушный транспорт. Система управления безопасностью вертолетной деятельности. Термины и определения [Электронный ресурс]. - URL: https://base.garant.ru/71396958/ - Дата обращения: 16.05.2026.",
]


def set_run_font(run, size=14):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_biblio_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)


def paragraph_insert_after(paragraph, text):
    new_p = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_p._p)
    format_biblio_paragraph(new_p)
    return new_p


def main():
    doc = Document(DOCX_PATH)
    texts = [p.text.strip() for p in doc.paragraphs]
    if any("ГОСТ Р 55585-2013" in t for t in texts):
        doc.save(DOCX_PATH)
        return

    # Locate bibliography block and boundary before appendix.
    start = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().lower().startswith("список литературы"):
            start = idx
            break
    if start is None:
        raise RuntimeError("Не найден раздел списка литературы")

    appendix_idx = None
    for idx in range(start + 1, len(doc.paragraphs)):
        text = doc.paragraphs[idx].text.strip().lower()
        if text.startswith("приложение"):
            appendix_idx = idx
            break
    if appendix_idx is None:
        appendix_idx = len(doc.paragraphs)

    # Use the last non-empty paragraph before appendix as insertion point.
    last_idx = None
    for idx in range(appendix_idx - 1, start, -1):
        if doc.paragraphs[idx].text.strip():
            last_idx = idx
            break
    if last_idx is None:
        raise RuntimeError("Не найдено место для вставки в списке литературы")

    # Try to continue explicit numeric numbering when present in plain text.
    number_re = re.compile(r"^(\d+)\.\s+")
    last_num = 21
    for idx in range(start + 1, appendix_idx):
        m = number_re.match(doc.paragraphs[idx].text.strip())
        if m:
            last_num = int(m.group(1))

    cursor = doc.paragraphs[last_idx]
    next_num = last_num + 1
    for item in GOST_ITEMS:
        cursor = paragraph_insert_after(cursor, f"{next_num}. {item}")
        next_num += 1

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
