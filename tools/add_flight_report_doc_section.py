from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/kotewa/my/diplom на проверку.docx")
REPORT_SCREENSHOT = Path("/Users/kotewa/my/скрины/Снимок экрана 2026-05-13 в 22.44.03.png")
CAPTION = "Рисунок 14 - Печатная форма отчета по рейсу"


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_after=0):
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.5 if first_indent else 0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        set_run_font(run)


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    new_p.addnext(result._p)
    paragraph._p.getparent().remove(new_p)
    if text:
        run = result.add_run(text)
        set_run_font(run)
    return result


def replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)
    format_paragraph(paragraph)


def contains_caption(doc: Document) -> bool:
    return any(p.text.strip() == CAPTION for p in doc.paragraphs)


def main():
    doc = Document(DOCX_PATH)
    if contains_caption(doc):
        return

    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Для выбранного рейса реализована печатная форма отчета"):
            anchor = paragraph
            break
    if anchor is None:
        raise RuntimeError("Report paragraph was not found")

    replace_paragraph_text(
        anchor,
        "Для выбранного рейса реализована отчетность в виде печатной формы, которая открывается из карточки рейса с помощью кнопки «Печать отчета». Отчет содержит основные сведения о рейсе, тип воздушного судна, маршрут, плановое время вылета и прилета, оценку рисков по этапам, ключевые метеорологические факторы, принятое решение диспетчера, историю изменений по рейсу и поля для подписей ответственных лиц. Такая форма позволяет не только просматривать результат анализа на экране, но и фиксировать принятое диспетчерское решение в виде отдельного документа.",
    )

    mention = paragraph_after(
        anchor,
        "Пример сформированного оперативного диспетчерского отчета по рейсу представлен на рисунке 14.",
    )
    format_paragraph(mention)

    pic_p = paragraph_after(mention)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.first_line_indent = Cm(0)
    pic_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pic_p.paragraph_format.space_before = Pt(6)
    pic_p.paragraph_format.space_after = Pt(0)
    run = pic_p.add_run()
    run.add_picture(str(REPORT_SCREENSHOT), width=Cm(11.5))

    caption = paragraph_after(pic_p, CAPTION)
    format_paragraph(caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=6)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
