from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/kotewa/my/diplom на проверку.docx")

SOURCES = [
    "ГОСТ 19.101-77. Единая система программной документации. Виды программ и программных документов. - М.: Стандартинформ, 1978.",
    "ГОСТ 19.201-78. Единая система программной документации. Техническое задание на программу. Требования к содержанию и оформлению. - М.: Стандартинформ, 1979.",
    "ГОСТ 19.404-79. Единая система программной документации. Описание программы. Требования к содержанию и оформлению. - М.: Стандартинформ, 1980.",
    "ГОСТ 34.201-89. Информационная технология. Виды, комплектность и обозначения документов при создании автоматизированных систем. - М.: Изд-во стандартов, 1990.",
    "ГОСТ 34.602-89. Информационная технология. Техническое задание на создание автоматизированной системы. - М.: Изд-во стандартов, 1991.",
    "ГОСТ Р 57240-2016. Воздушный транспорт. Менеджмент безопасности авиационной деятельности в гражданской авиации. Основные положения [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/6094796d-8ed9-48b0-b222-09e48ffc05f0 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 57908-2017. Воздушный транспорт. Система менеджмента безопасности авиационной деятельности. База данных. Авиационные риски по реализации системы оценки безопасности полетов при обеспечении воздушного движения [Электронный ресурс]. - URL: https://internet-law.ru/gosts/gost/65810/ - Дата обращения: 16.05.2026.",
    "ГОСТ Р 58712-2019. Автоматизированная метеорологическая измерительная система. Общие технические требования [Электронный ресурс]. - URL: https://www.standards.ru/print.aspx?control=27&id=8096369&print=yes - Дата обращения: 16.05.2026.",
    "ГОСТ 22.1.01-2023. Безопасность в чрезвычайных ситуациях. Мониторинг и прогнозирование. Основные положения [Электронный ресурс]. - URL: https://www.gostinfo.ru/catalog/Details/?id=7476806 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 22.1.07-2023. Безопасность в чрезвычайных ситуациях. Мониторинг и прогнозирование опасных метеорологических явлений и процессов. Общие требования [Электронный ресурс]. - URL: https://base.garant.ru/408555345/ - Дата обращения: 16.05.2026.",
    "ГОСТ Р 56079-2014. Изделия авиационной техники. Безопасность полета, надежность, контролепригодность, эксплуатационная и ремонтная технологичность. Номенклатура показателей [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/fc9bc901-8f2c-40f5-9ada-b100881a43c5 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 55585-2013. Воздушный транспорт. Система управления безопасностью полетов воздушных судов. Термины и определения [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/9e224af2-9c65-4ccc-9d35-fa2dcc116ef5 - Дата обращения: 16.05.2026.",
    "Богаткин О. Г. Авиационная метеорология: учебник. - СПб.: Изд-во РГГМУ, 2005. - 328 с.",
    "Блинов И. В., Соловьев А. А. Основы проектирования информационных систем. - М.: Академия, 2020. - 320 с.",
    "Воронин Е. Н. Введение в проектирование информационных систем. - М.: ФИЗМАТЛИТ, 2021. - 256 с.",
    "Горшков И. А. Проектирование баз данных: учебное пособие. - М.: Юрайт, 2022. - 314 с.",
    "Ермаков С. Ю. Информационные технологии в управлении бизнес-процессами. - М.: Инфра-М, 2021. - 368 с.",
    "Соловьев А. А., Терехов А. Н. Информационные системы и технологии. - М.: Горячая линия - Телеком, 2019. - 480 с.",
    "Spring. Spring Boot Reference Documentation [Электронный ресурс]. - URL: https://docs.spring.io/spring-boot/ - Дата обращения: 12.01.2026.",
    "Oracle. Java Platform, Standard Edition Documentation [Электронный ресурс]. - URL: https://docs.oracle.com/en/java/ - Дата обращения: 13.01.2026.",
    "React. React Documentation [Электронный ресурс]. - URL: https://react.dev/ - Дата обращения: 20.01.2026.",
    "PostgreSQL Global Development Group. PostgreSQL Documentation [Электронный ресурс]. - URL: https://www.postgresql.org/docs/ - Дата обращения: 12.03.2026.",
    "Docker Inc. Docker Documentation [Электронный ресурс]. - URL: https://docs.docker.com/ - Дата обращения: 25.04.2026.",
    "OpenWeatherMap. Weather API [Электронный ресурс]. - URL: https://openweathermap.org/api - Дата обращения: 20.02.2026.",
    "MDN Web Docs. Web Storage API [Электронный ресурс]. - URL: https://developer.mozilla.org/en-US/docs/Web/API/Web_Storage_API - Дата обращения: 05.05.2026.",
    "Leaflet. Leaflet Documentation [Электронный ресурс]. - URL: https://leafletjs.com/reference.html - Дата обращения: 25.02.2026.",
    "OpenStreetMap Foundation. OpenStreetMap Wiki [Электронный ресурс]. - URL: https://wiki.openstreetmap.org/ - Дата обращения: 18.03.2026.",
]


def set_run_font(run, size=14):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_biblio(paragraph):
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    ppr = paragraph._element.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)
    for run in paragraph.runs:
        set_run_font(run)


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    new_p.addnext(result._p)
    paragraph._p.getparent().remove(new_p)
    if text:
        result.add_run(text)
    return result


def main():
    doc = Document(DOCX_PATH)
    biblio_idx = None
    appendix_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().lower()
        if biblio_idx is None and t.startswith("список литературы"):
            biblio_idx = i
        elif biblio_idx is not None and t.startswith("приложение"):
            appendix_idx = i
            break
    if biblio_idx is None or appendix_idx is None:
        raise RuntimeError("Не найден раздел литературы или Приложение")

    for p in list(doc.paragraphs[biblio_idx + 1 : appendix_idx]):
        p._element.getparent().remove(p._element)

    anchor = doc.paragraphs[biblio_idx]
    for i, src in enumerate(SOURCES, start=1):
        anchor = paragraph_after(anchor, f"{i}. {src}")
        format_biblio(anchor)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
