from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


DOCX_PATH = Path("/Users/kotewa/my/diplom.docx")
SCREENS = Path("/Users/kotewa/my/скрины")


SCREEN_ITEMS = [
    {
        "anchor": "В клиентской части реализован выбор роли пользователя",
        "mention": "Экран выбора роли пользователя показан на рисунке 6. Он позволяет запустить демонстрационный сценарий от лица диспетчера или метеоролога.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.29.51.png",
        "caption": "Рисунок 6 - Экран выбора роли пользователя",
        "width": 15.0,
    },
    {
        "anchor": "Пользовательский интерфейс реализует два режима работы",
        "mention": "Главный экран диспетчера представлен на рисунке 7. На нем объединены сведения о рейсе, погодные параметры, карта маршрута и область рекомендаций.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.23.59.png",
        "caption": "Рисунок 7 - Главный экран диспетчера",
        "width": 16.0,
    },
    {
        "anchor": "Форма создания рейса содержит поля для ввода номера рейса",
        "mention": "Форма создания рейса с заполненными данными показана на рисунке 8. Пользователь выбирает города и аэропорты из справочника, указывает дату, код авиакомпании и номер рейса.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.25.55.png",
        "caption": "Рисунок 8 - Форма создания рейса с выбором аэропортов из справочника",
        "width": 15.0,
    },
    {
        "anchor": "Карта маршрута построена на Leaflet и OpenStreetMap",
        "mention": "Пример отображения построенного маршрута на карте приведен на рисунке 9. На карте показаны аэропорты вылета и прибытия, а также линия маршрута между ними.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.24.33.png",
        "caption": "Рисунок 9 - Результат построения маршрута на карте",
        "width": 16.0,
    },
    {
        "anchor": "Отдельное внимание уделяется визуальному отображению риска",
        "mention": "Отображение оценки риска по погодным параметрам показано на рисунке 10. Пользователь видит частные показатели риска, итоговый балл и причины повышения риска.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.25.32.png",
        "caption": "Рисунок 10 - Отображение оценки риска по погодным параметрам",
        "width": 12.0,
    },
    {
        "anchor": "Список рейсов позволяет диспетчеру видеть несколько рейсов одновременно",
        "mention": "Список рейсов с фильтрами и статусами показан на рисунке 11. Такой экран помогает диспетчеру быстро сравнивать рейсы по времени, маршруту, уровню риска и принятому решению.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.27.11.png",
        "caption": "Рисунок 11 - Список рейсов с фильтрами и статусами",
        "width": 16.0,
    },
    {
        "anchor": "Кнопка «Запросить данные у метеоролога» отображается",
        "mention": "Интерфейс формирования запроса метеорологу приведен на рисунке 12. В форме фиксируются рейс, причина запроса и набор метеорологических параметров, требующих уточнения.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.28.40.png",
        "caption": "Рисунок 12 - Интерфейс формирования запроса метеорологу",
        "width": 16.0,
    },
    {
        "anchor": "В списке рейсов реализованы фильтры по уровню риска",
        "mention": "Панель аналитики рейсов показана на рисунке 13. Она используется для обобщенной оценки количества рейсов, распределения рисков и активности запросов к метеорологу.",
        "image": SCREENS / "Снимок экрана 2026-05-13 в 21.29.11.png",
        "caption": "Рисунок 13 - Панель аналитики рейсов",
        "width": 16.0,
    },
]


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_after=0):
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.5 if first_indent else 0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        set_run_font(run)


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    new_p.addnext(result._p)
    paragraph._p.getparent().remove(new_p)
    if text:
        run = result.add_run(text)
        set_run_font(run)
    return result


def insert_after(paragraph, text, image_path: Path, caption: str, width_cm: float):
    mention = paragraph_after(paragraph, text)
    format_paragraph(mention)

    pic_p = paragraph_after(mention)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.first_line_indent = Cm(0)
    pic_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pic_p.paragraph_format.space_before = Pt(6)
    pic_p.paragraph_format.space_after = Pt(0)
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_p = paragraph_after(pic_p, caption)
    format_paragraph(caption_p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=6)


def remove_old_appended_screens(doc: Document):
    start_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("Рисунок 6 - Главный экран диспетчера"):
            start_index = max(0, index - 2)
            break
    if start_index is None:
        return
    for paragraph in list(doc.paragraphs[start_index:]):
        paragraph._element.getparent().remove(paragraph._element)


def contains_caption(doc: Document, caption: str) -> bool:
    return any(paragraph.text.strip() == caption for paragraph in doc.paragraphs)


def main():
    doc = Document(DOCX_PATH)
    remove_old_appended_screens(doc)

    for item in SCREEN_ITEMS:
        if contains_caption(doc, item["caption"]):
            continue
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith(item["anchor"]):
                width = item["width"]
                with Image.open(item["image"]) as image:
                    if image.height > image.width * 1.15:
                        width = min(width, 10.0)
                insert_after(paragraph, item["mention"], item["image"], item["caption"], width)
                break
        else:
            raise RuntimeError(f"Anchor not found: {item['anchor']}")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
