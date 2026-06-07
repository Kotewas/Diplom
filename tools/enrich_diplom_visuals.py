from __future__ import annotations

import html
import math
import re
import textwrap
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


DOCX_PATH = Path("/Users/kotewa/my/diplom.docx")
DRAWIO_PATH = Path("/Users/kotewa/my/Диплом схемы/use-case.drawio")
USE_CASE_PNG = Path("/private/tmp/diplom_use_case.png")

SCREENSHOTS = [
    (
        Path("/Users/kotewa/my/скрины/Снимок экрана 2026-04-25 в 18.40.15.png"),
        "Рисунок 6 - Главный экран диспетчера с формой создания рейса и картой маршрута",
    ),
    (
        Path("/Users/kotewa/my/скрины/Снимок экрана 2026-05-05 в 20.59.37.png"),
        "Рисунок 7 - Отображение аэропортов и маршрутов на интерактивной карте",
    ),
    (
        Path("/Users/kotewa/my/скрины/Снимок экрана 2026-05-05 в 21.04.33.png"),
        "Рисунок 8 - Отображение погодных параметров и уровня риска в интерфейсе",
    ),
]


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def clean_label(value: str) -> str:
    text = re.sub(r"<[^>]+>", " ", html.unescape(value or ""))
    return " ".join(text.split())


def draw_wrapped_text(draw: ImageDraw.ImageDraw, box, text: str, font, fill="black", align="center"):
    x, y, w, h = box
    max_chars = max(8, int(w / 8))
    lines = []
    for part in text.split("\n"):
        lines.extend(textwrap.wrap(part, width=max_chars) or [""])
    line_h = font.getbbox("Ag")[3] - font.getbbox("Ag")[1] + 4
    total_h = line_h * len(lines)
    cy = y + (h - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        if align == "center":
            tx = x + (w - (bbox[2] - bbox[0])) / 2
        else:
            tx = x
        draw.text((tx, cy), line, font=font, fill=fill)
        cy += line_h


def draw_actor(draw: ImageDraw.ImageDraw, cx: float, y: float, label: str, font):
    head_r = 18
    draw.ellipse((cx - head_r, y, cx + head_r, y + head_r * 2), outline="black", width=3)
    body_top = y + head_r * 2
    body_bottom = body_top + 55
    draw.line((cx, body_top, cx, body_bottom), fill="black", width=3)
    draw.line((cx - 34, body_top + 18, cx + 34, body_top + 18), fill="black", width=3)
    draw.line((cx, body_bottom, cx - 28, body_bottom + 42), fill="black", width=3)
    draw.line((cx, body_bottom, cx + 28, body_bottom + 42), fill="black", width=3)
    draw_wrapped_text(draw, (cx - 70, body_bottom + 48, 140, 40), label, font)


def create_use_case_png():
    root = ET.parse(DRAWIO_PATH).getroot()
    cells = {}
    edges = []
    for cell in root.iter("mxCell"):
        geom = cell.find("mxGeometry")
        if geom is None:
            continue
        cell_id = cell.attrib.get("id")
        if not cell_id:
            continue
        item = {
            "id": cell_id,
            "value": clean_label(cell.attrib.get("value", "")),
            "style": cell.attrib.get("style", ""),
            "x": float(geom.attrib.get("x", 0)),
            "y": float(geom.attrib.get("y", 0)),
            "w": float(geom.attrib.get("width", 0)),
            "h": float(geom.attrib.get("height", 0)),
        }
        if cell.attrib.get("edge") == "1":
            item["source"] = cell.attrib.get("source")
            item["target"] = cell.attrib.get("target")
            edges.append(item)
        else:
            cells[cell_id] = item

    visible = [item for item in cells.values() if item["value"] or "umlActor" in item["style"] or item["w"] > 300]
    min_x = min(item["x"] for item in visible)
    min_y = min(item["y"] for item in visible)
    max_x = max(item["x"] + item["w"] for item in visible)
    max_y = max(item["y"] + item["h"] for item in visible)
    margin = 70
    scale = min(1.25, 1450 / (max_x - min_x + margin * 2), 1750 / (max_y - min_y + margin * 2))
    width = int((max_x - min_x) * scale + margin * 2)
    height = int((max_y - min_y) * scale + margin * 2)

    def tr(x, y):
        return ((x - min_x) * scale + margin, (y - min_y) * scale + margin)

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = load_font(24)
    small_font = load_font(21)
    title_font = load_font(28)

    for item in visible:
        if item["w"] > 300 and not item["value"]:
            x, y = tr(item["x"], item["y"])
            draw.rectangle((x, y, x + item["w"] * scale, y + item["h"] * scale), outline="black", width=3)
            draw.text((x + item["w"] * scale / 2 - 55, y + 12), "Система", font=title_font, fill="black")

    for edge in edges:
        source = cells.get(edge.get("source"))
        target = cells.get(edge.get("target"))
        if not source or not target:
            continue
        sx, sy = tr(source["x"] + source["w"] / 2, source["y"] + source["h"] / 2)
        tx, ty = tr(target["x"] + target["w"] / 2, target["y"] + target["h"] / 2)
        draw.line((sx, sy, tx, ty), fill=(80, 80, 80), width=2)

    for item in visible:
        x, y = tr(item["x"], item["y"])
        w = item["w"] * scale
        h = item["h"] * scale
        if "umlActor" in item["style"]:
            draw_actor(draw, x + w / 2, y, item["value"], font)
        elif "ellipse" in item["style"]:
            draw.ellipse((x, y, x + w, y + h), outline="black", width=3, fill=(248, 252, 255))
            draw_wrapped_text(draw, (x + 10, y + 8, w - 20, h - 16), item["value"], small_font)

    img.save(USE_CASE_PNG)


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def qn(tag: str):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def format_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.5)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)


def insert_paragraph_after(paragraph, text: str):
    new_p = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_p._p)
    format_body(new_p)
    return new_p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run)


def add_picture_with_caption(doc: Document, path: Path, caption: str, width_cm: float = 16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def contains_text(doc: Document, needle: str) -> bool:
    return any(needle in paragraph.text for paragraph in doc.paragraphs)


def enrich_docx():
    create_use_case_png()
    doc = Document(DOCX_PATH)

    if not contains_text(doc, "Диаграмма вариантов использования представлена"):
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith("Функциональная схема разработанной системы представлена"):
                insert_paragraph_after(
                    paragraph,
                    "Диаграмма вариантов использования представлена в приложении на рисунке 5. Она показывает двух основных участников системы: диспетчера и метеоролога. Диспетчер создает рейс, просматривает список рейсов, анализирует метеорологические данные, риск и рекомендации, формирует запрос на уточнение информации и принимает итоговое решение по рейсу. Метеоролог получает запрос от диспетчера и передает уточненные метеорологические данные для дальнейшего анализа.",
                )
                break

    if not contains_text(doc, "В приложении также приведены изображения пользовательского интерфейса"):
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith("Пользовательский интерфейс реализует два режима работы"):
                insert_paragraph_after(
                    paragraph,
                    "В приложении также приведены изображения пользовательского интерфейса. На них показаны главный экран диспетчера, интерактивная карта с аэропортами и маршрутом, а также отображение погодных параметров и уровня риска. Эти материалы подтверждают, что разработанная система имеет не только серверную логику, но и готовый пользовательский сценарий работы.",
                )
                break

    if contains_text(doc, "Рисунок 5 - Диаграмма вариантов использования системы"):
        doc.save(DOCX_PATH)
        return

    doc.add_page_break()
    add_picture_with_caption(doc, USE_CASE_PNG, "Рисунок 5 - Диаграмма вариантов использования системы", 16.0)
    for screenshot, caption in SCREENSHOTS:
        if screenshot.exists():
            doc.add_page_break()
            width = 16.0
            with Image.open(screenshot) as image:
                if image.height > image.width * 1.15:
                    width = 9.5
            add_picture_with_caption(doc, screenshot, caption, width)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    enrich_docx()
