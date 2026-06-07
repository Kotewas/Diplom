from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "diplom.before-fix.docx"
TARGET = ROOT / "diplom.docx"

PROJECT_TITLE = (
    "Разработка веб-системы мониторинга и анализа погодных условий "
    "для авиационных рейсов"
)
CODE = "ДП.09.02.07.66.25.00.00 ПЗ"
SCHEME_DIR = Path("/Users/kotewa/my/Диплом схемы")
SCHEMES = [
    (
        SCHEME_DIR / "Функциональная схема.png",
        "Рисунок В.1 - Функциональная схема веб-системы мониторинга погодных условий",
        "Функциональная схема системы",
        Cm(16.0),
        None,
    ),
    (
        SCHEME_DIR / "BPMN Создание рейса.png",
        "Рисунок В.2 - BPMN-схема процесса создания рейса",
        "BPMN-схема создания рейса",
        Cm(17.0),
        None,
    ),
    (
        SCHEME_DIR / "BPMN запрос метеорологу.png",
        "Рисунок В.3 - BPMN-схема запроса данных у метеоролога",
        "BPMN-схема запроса метеорологу",
        Cm(17.0),
        None,
    ),
    (
        SCHEME_DIR / "Перерасчет риска.png",
        "Рисунок В.4 - Схема процесса перерасчета риска рейса",
        "Схема перерасчета риска",
        None,
        Cm(22.0),
    ),
]


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def set_cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def append_cell_paragraph(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=12)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 3"].font.size = Pt(12)

    if "Diplom Caption" not in styles:
        cap = styles.add_style("Diplom Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        cap.font.size = Pt(12)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        cap.paragraph_format.space_after = Pt(0)


def add_centered(doc, text="", size=12, bold=False, spacing_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_plain(doc, text="", first_indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.5 if first_indent else 0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_plain_runs(doc, parts, first_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.5 if first_indent else 0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, bold=bold)
    return p


def add_formula(doc, expression, number, explanations):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = table.rows[0].cells
    set_cell_width(cells[0], 15.5)
    set_cell_width(cells[1], 1.5)

    formula_p = cells[0].paragraphs[0]
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_p.paragraph_format.first_line_indent = Cm(0)
    formula_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    formula_p.paragraph_format.space_after = Pt(0)
    formula_run = formula_p.add_run(expression)
    set_run_font(formula_run)

    number_p = cells[1].paragraphs[0]
    number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_p.paragraph_format.first_line_indent = Cm(0)
    number_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    number_p.paragraph_format.space_after = Pt(0)
    number_run = number_p.add_run(f"({number})")
    set_run_font(number_run)

    for index, explanation in enumerate(explanations):
        prefix = "где " if index == 0 else ""
        add_plain(doc, prefix + explanation, first_indent=False)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_picture_centered(doc, image_path, width=None, height=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    if width is not None and height is not None:
        run.add_picture(str(image_path), width=width, height=height)
    elif width is not None:
        run.add_picture(str(image_path), width=width)
    else:
        run.add_picture(str(image_path), height=height)
    return p


def add_title_page(doc):
    add_centered(doc, "МИНИСТЕРСТВО ПРОМЫШЛЕННОСТИ И ТОРГОВЛИ ТВЕРСКОЙ ОБЛАСТИ", 12, True)
    add_centered(doc, "ГБПОУ «Тверской колледж им. А.Н.Коняева»", 12, True)
    for _ in range(6):
        add_centered(doc)
    add_centered(doc, "ДИПЛОМНЫЙ ПРОЕКТ", 16, True)
    add_centered(doc)
    add_centered(doc, f"На тему: «{PROJECT_TITLE}»", 12, False)
    for _ in range(3):
        add_centered(doc)
    add_plain(doc, "Студента группы ________________________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Ф.И.О. _________________________________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Специальность 09.02.07 «Информационные системы и программирование»", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Руководитель ___________________________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Рецензент ______________________________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    for _ in range(3):
        add_centered(doc)
    add_plain(doc, "К защите допущен(а)", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "заместитель директора по УР ____________________ Н.С.Лукина", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    for _ in range(5):
        add_centered(doc)
    add_centered(doc, "Тверь")
    add_centered(doc, "2026")


def add_assignment(doc):
    doc.add_page_break()
    add_centered(doc, "МИНИСТЕРСТВО ПРОМЫШЛЕННОСТИ И ТОРГОВЛИ ТВЕРСКОЙ ОБЛАСТИ", 12, True)
    add_centered(doc, "ГБПОУ «Тверской колледж им. А.Н.Коняева»", 12, True)
    add_plain(doc, "УТВЕРЖДАЮ", first_indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_plain(doc, "Заведующий отделением _____________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_plain(doc, "«___» __________ 2026 г.", first_indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_centered(doc)
    add_centered(doc, "ДИПЛОМНОЕ ЗАДАНИЕ", 14, True)
    add_plain(doc, "Студенту(ке) ________________________________________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, f"Тема: «{PROJECT_TITLE}».", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Содержание расчетно-пояснительной записки", first_indent=False, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in [
        "Введение",
        "1 Теоретический раздел",
        "1.1 Особенности диспетчерского управления авиационными рейсами",
        "1.2 Метеофакторы, влияющие на выполнение рейсов",
        "1.3 Методы оценки риска и принятия решений в авиации",
        "1.4 Технологии разработки информационных систем",
        "1.5 Сравнительный анализ существующих решений мониторинга и анализа погоды",
        "2 Практический раздел",
        "2.1 Обоснование необходимости разработки системы и постановка цели",
        "2.2 Техническое задание",
        "2.3 Архитектура системы",
        "2.4 Разработка серверной части",
        "2.5 Интеграция с погодным API и обработка метеоданных",
        "2.6 Реализация алгоритма расчета уровня риска рейса",
        "2.7 Разработка клиентской части",
        "2.8 Проектирование и реализация базы данных",
        "2.9 Контейнеризация и развертывание",
        "2.10 Тестирование и верификация корректности работы системы",
        "2.11 Оценка результатов внедрения",
        "3 Обеспечение безопасности и надежности данных",
        "Заключение",
        "Список литературы и источников",
        "Приложение А Структура проекта",
        "Приложение Б Инструкция по запуску проекта",
        "Приложение В Графические схемы проекта",
    ]:
        add_plain(doc, item, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_centered(doc)
    add_plain(doc, "Срок окончания выполнения работы ___________________ 2026 г.", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Руководитель дипломной работы ______________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Дата выдачи задания «___» __________________ 2026 г.", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_plain(doc, "Студент ___________________        _____________________________", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)


CONTENTS = [
    ("Введение", 4),
    ("1 Теоретический раздел", 7),
    ("1.1 Особенности диспетчерского управления авиационными рейсами", 7),
    ("1.2 Метеофакторы, влияющие на выполнение рейсов", 9),
    ("1.3 Методы оценки риска и принятия решений в авиации", 11),
    ("1.4 Технологии разработки информационных систем", 13),
    ("1.5 Сравнительный анализ существующих решений мониторинга и анализа погоды", 15),
    ("2 Практический раздел", 17),
    ("2.1 Обоснование необходимости разработки системы и постановка цели", 17),
    ("2.2 Техническое задание", 18),
    ("2.3 Архитектура системы", 21),
    ("2.4 Разработка серверной части", 23),
    ("2.5 Интеграция с погодным API и обработка метеоданных", 25),
    ("2.6 Реализация алгоритма расчета уровня риска рейса", 27),
    ("2.7 Разработка клиентской части", 30),
    ("2.8 Проектирование и реализация базы данных", 32),
    ("2.9 Контейнеризация и развертывание", 34),
    ("2.10 Тестирование и верификация корректности работы системы", 35),
    ("2.11 Оценка результатов внедрения", 38),
    ("3 Обеспечение безопасности и надежности данных", 40),
    ("3.1 Анализ рисков безопасности веб-приложения", 40),
    ("3.2 Реализация защиты данных", 42),
    ("3.3 Обеспечение надежности и отказоустойчивости", 44),
    ("Заключение", 46),
    ("Список литературы и источников", 49),
    ("Приложение А Структура проекта", 51),
    ("Приложение Б Инструкция по запуску проекта", 52),
    ("Приложение В Графические схемы проекта", 53),
]


def add_contents(doc):
    doc.add_page_break()
    add_centered(doc, "Содержание", 14, True)
    for title, page in CONTENTS:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(17), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run)


def extract_clean_body():
    doc = Document(SOURCE)
    paragraphs = []
    skip = False
    for p in doc.paragraphs:
        text = p.text.replace("\xa0", " ")
        text = re.sub(r"[ \t\r\n]+", " ", text).strip()
        if not text:
            continue
        if text.startswith("Список литературы"):
            break
        if text in {"Изм.", "Лист", "№ докум.", "Подпись", "Дата", "Дат", "П"}:
            continue
        if re.fullmatch(r"\d+", text):
            continue
        if CODE in text:
            continue
        text = text.replace("REQUEST_WEATHER_UPDATE", "REQUEST_METEO")
        text = text.replace(
            "использование последнего валидного набора данных с отметкой об устаревании либо уточнение информации через имитацию взаимодействия с метеорологом",
            "использование свежего локального кэша, синтетического fallback-набора при сбое внешнего API либо уточнение информации через имитацию взаимодействия с метеорологом",
        )
        text = text.replace(
            "Кэширование используется для хранения последнего валидного набора погодных данных. Если внешний API временно недоступен, система может показать последние полученные данные, но обязательно должна отметить их как устаревшие. Это важно, чтобы диспетчер понимал ограниченность использования такой информации.",
            "Кэширование используется для временного хранения свежих погодных данных и снижения количества обращений к внешнему API. Если внешний API временно недоступен и свежего кэша нет, серверная часть формирует синтетический fallback-набор с признаком provider = synthetic-fallback, а интерфейс отдельно показывает, что данные получены по резервной модели.",
        )
        text = text.replace(
            "В ходе тестирования была подтверждена корректность работы основных функций системы, включая обработку различных сценариев: нормальную работу с погодными данными, сбои внешнего API, использование кэшированных данных, перерасчет риска и работу пользовательского сценария взаимодействия диспетчера с метеорологом.",
            "В ходе функциональной проверки была подтверждена работоспособность основных сценариев системы: создание рейса, получение погодных данных, резервная обработка сбоя внешнего API через fallback-модель, перерасчет риска, моделирование задержки и пользовательский сценарий взаимодействия диспетчера с метеорологом.",
        )
        text = text.replace(
            "использование кэшированных данных",
            "использование свежего кэша и резервной fallback-модели",
        )
        text = text.replace(
            "использование последнего валидного набора погодных данных при сбое API",
            "использование свежего кэша и резервной fallback-модели при сбое API",
        )
        text = text.replace(
            "Таблица — Использование последнего валидного набора данных при сбое API",
            "Таблица 2 - Проверка резервной fallback-модели при сбое API",
        )
        text = text.replace(
            "Таблица 1 - Проверка резервной fallback-модели при сбое API",
            "Таблица 5 - Проверка резервной fallback-модели при сбое API",
        )
        text = text.replace(
            "Таблица 2 - Проверка резервной fallback-модели при сбое API",
            "Таблица 5 - Проверка резервной fallback-модели при сбое API",
        )
        text = text.replace(
            "последние данные: ветер 5 м/с, видимость 10 км, температура +7°C",
            "резервные данные: ветер 5 м/с, видимость 10 км, температура +7°C",
        )
        text = text.replace(
            "Последние данные отображены, предупреждение показано",
            "При недоступности API система отображает резервные данные с признаком fallback, интерфейс показывает предупреждение о резервном источнике",
        )
        text = text.replace(
            "Система отображает последние валидные данные, помечает их как устаревшие и предлагает запрос к метеослужбе",
            "Система отображает резервные данные с признаком fallback и предлагает запрос к метеослужбе",
        )
        text = text.replace(
            "Данные устарели",
            "Данные получены из резервного источника",
        )
        if text.startswith("Универсальные погодные сервисы предоставляют информацию о погоде"):
            text = (
                "Универсальные погодные API предоставляют информацию о погоде для широкого круга пользователей. "
                "В разработанной системе используется OpenWeatherMap API, который позволяет получать данные о "
                "температуре, ветре, осадках, давлении и видимости. Его преимуществами являются доступность, "
                "простота интеграции и возможность использования в учебных и демонстрационных проектах."
            )
        paragraphs.append(text)
    return paragraphs


EXTRA_BY_HEADING = {
    "Введение": [
        "В ходе проектирования были подготовлены графические материалы: функциональная схема системы, BPMN-схема создания рейса, BPMN-схема запроса данных у метеоролога и схема процесса перерасчета риска. Указанные схемы используются как проектная графическая часть и приведены в приложении В.",
    ],
    "2.3 Архитектура системы": [
        "В текущей версии проекта архитектура дополнена модулем истории изменений рейса. После ручного обновления риска или применения решения диспетчера backend сохраняет запись в таблице flight_history, где фиксируются идентификатор рейса, время изменения, прежний и новый уровень риска, погодный снимок, принятое решение, причина решения и величина задержки при наличии. Это позволяет не только видеть актуальное состояние рейса, но и анализировать последовательность действий диспетчера.",
        "Серверная часть предоставляет отдельные REST-эндпоинты для получения истории рейса, ручного пересчета риска, моделирования задержки, применения решения диспетчера и удаления рейса. Клиентская часть использует эти эндпоинты через слой сервисов flightsApi.js, что отделяет UI-компоненты от деталей HTTP-запросов и обработки ошибок.",
        "Функциональная схема разработанной системы представлена в приложении В на рисунке В.1. Она отражает основные компоненты приложения: ввод данных рейса, получение погодной информации, расчет риска, взаимодействие диспетчера с метеорологом, сохранение данных и отображение результата пользователю.",
    ],
    "2.4 Разработка серверной части": [
        "В серверной части реализована поддержка двух типов воздушных судов: самолет и вертолет. Для самолетов используется формат номера рейса IATA, например SU123, а для вертолетов применяется формат H и числовой идентификатор, например H123. Тип воздушного судна влияет на расчет планового времени прибытия: для самолета используется более высокая крейсерская скорость и больший операционный запас времени, для вертолета - меньшая скорость и меньший минимальный интервал выполнения рейса.",
        "Для вертолетных сценариев справочник аэропортов расширен набором площадок Тверской области. В него включены площадки Тверь, Торжок, Ржев, Конаково, Кашин, Бежецк, Бологое, Вышний Волочек, Кимры, Калязин, Осташков и другие точки. Это позволяет демонстрировать работу системы не только на межрегиональных авиационных маршрутах, но и на региональных вертолетных рейсах.",
        "Также реализовано автоматическое первичное решение по рейсу. При создании карточки система может оставить решение в статусе ожидания, автоматически рекомендовать выполнение рейса при низком риске или предложить задержку при повышенном, но не критическом риске. Окончательное решение все равно остается за диспетчером и может быть зафиксировано вручную с обязательным указанием причины.",
    ],
    "2.5 Интеграция с погодным API и обработка метеоданных": [
        "Погодные данные в системе запрашиваются через OpenWeatherMap API. Backend обращается к эндпоинту openweathermap.org, получает текущие метеорологические параметры по координатам аэропорта и передает нормализованный результат клиентской части. Если внешний сервис временно недоступен или ключ API не задан в локальной среде, система использует только внутренний синтетический fallback-набор, чтобы интерфейс оставался работоспособным в демонстрационном режиме без обращения к другим погодным API.",
        "В интерфейсе данные, полученные из fallback-источника, визуально помечаются. Это важно для прозрачности: пользователь видит, что расчет выполнен не по основному источнику, а по резервной модели, и может дополнительно запросить уточнение у метеоролога.",
    ],
    "2.6 Реализация алгоритма расчета уровня риска рейса": [
        "Расчет риска выполняется по балльной модели. Каждый неблагоприятный фактор добавляет к частному риску заданное количество баллов, после чего значение ограничивается диапазоном от 0 до 100. Для ограничения результата используется формула (1).",
        {
            "type": "formula",
            "number": 1,
            "expression": "F(x) = min(100; max(0; round(x)))",
            "explanations": [
                "F(x) - значение риска после округления и ограничения, балл;",
                "x - расчетное значение риска до ограничения, балл.",
            ],
        },
        "Для аэропортов вылета и прибытия используется одинаковая формула поверхностного риска (2).",
        {
            "type": "formula",
            "number": 2,
            "expression": "Rпов = F(W + G + V + P + T + N + O + M)",
            "explanations": [
                "Rпов - поверхностный риск аэропорта, балл;",
                "W - вклад скорости ветра, балл;",
                "G - вклад порывов ветра, балл;",
                "V - вклад видимости, балл;",
                "P - вклад атмосферного давления, балл;",
                "T - вклад температуры воздуха, балл;",
                "N - вклад облачности, балл;",
                "O - вклад осадков, балл;",
                "M - вклад кода погодного явления, балл.",
            ],
        },
        "Маршрутный риск рассчитывается по формуле (3).",
        {
            "type": "formula",
            "number": 3,
            "expression": "Rмарш = F(10 + D + L + WV + PD)",
            "explanations": [
                "Rмарш - риск маршрута, балл;",
                "10 - базовое значение маршрутного риска, балл;",
                "D - вклад дальности маршрута, балл;",
                "L - вклад высокой широты маршрута, балл;",
                "WV - вклад ветрового фона в аэропортах вылета и прибытия, балл;",
                "PD - вклад барического контраста между аэропортами, балл.",
            ],
        },
        "Итоговый риск рейса определяется как взвешенная сумма частных рисков по формуле (4).",
        {
            "type": "formula",
            "number": 4,
            "expression": "Rитог = F(0,4 · Rвыл + 0,4 · Rприл + 0,2 · Rмарш)",
            "explanations": [
                "Rитог - итоговый риск рейса, балл;",
                "0,4 - весовой коэффициент риска аэропорта вылета;",
                "Rвыл - риск аэропорта вылета, балл;",
                "0,4 - весовой коэффициент риска аэропорта прибытия;",
                "Rприл - риск аэропорта прибытия, балл;",
                "0,2 - весовой коэффициент риска маршрута;",
                "Rмарш - риск маршрута, балл.",
            ],
        },
        "При моделировании задержки рассчитывается изменение риска с учетом продолжительности задержки и количества опасных сигналов по формуле (5).",
        {
            "type": "formula",
            "number": 5,
            "expression": "Rзад = F(Rтек - min(24; t · 0,22) · (0,25 + Q · 0,18))",
            "explanations": [
                "Rзад - риск после задержки рейса, балл;",
                "Rтек - текущий риск рейса, балл;",
                "24 - максимальное снижение риска при моделировании задержки, балл;",
                "t - продолжительность задержки, мин;",
                "0,22 - коэффициент снижения риска на одну минуту задержки;",
                "0,25 - базовый коэффициент влияния задержки;",
                "Q - количество опасных погодных сигналов;",
                "0,18 - коэффициент усиления влияния задержки при наличии опасных сигналов.",
            ],
        },
        "Если опасные сигналы отсутствуют, а текущий риск ниже 45 баллов, задержка может не снижать риск, а незначительно увеличивать его из-за смещения времени выполнения рейса.",
        "Для отображения результата применяется шкала: 0-29 баллов - низкий риск, 30-59 баллов - средний риск, 60-79 баллов - высокий риск, 80-100 баллов - критический риск. Эта шкала используется в интерфейсе для цветовой индикации карточек рейсов и принятия диспетчерского решения.",
        "В клиентской части дополнительно учитывается временной фактор рейса. Если до вылета остается слишком мало времени, система повышает риск из-за короткого окна подготовки. Если вылет запланирован далеко вперед, риск также может быть увеличен из-за меньшей устойчивости прогноза. Для ночных и ранних утренних вылетов применяется отдельная корректировка, особенно заметная для вертолетных рейсов.",
        "Интерфейс поддерживает сценарии чрезвычайной ситуации. Диспетчер может включить демонстрационный сценарий ЧС, после чего отображаемый риск корректируется на величину дополнительного штрафа. Это не заменяет серверный расчет, но позволяет показать, как система может использоваться при усложнении оперативной обстановки.",
        "При ответе метеоролога система оценивает полноту предоставленных данных. Если заполнены не все требуемые поля, клиентская логика может повысить коэффициент риска: при отсутствии данных сильнее, при частичном заполнении умеренно. Такой подход показывает, что неполная метеорологическая информация сама по себе является фактором неопределенности.",
    ],
    "2.7 Разработка клиентской части": [
        "Пользовательский интерфейс реализует два режима работы: режим диспетчера и режим метеоролога. Сессия пользователя сохраняется в localStorage, поэтому после обновления страницы выбранный сценарий работы может быть восстановлен. Диспетчер видит мониторинг рейсов, карту, список активных рейсов, рекомендации и форму создания рейса, а метеоролог - очередь запросов и форму ответа по выбранным метеорологическим параметрам.",
        "Карта маршрута построена на Leaflet и OpenStreetMap. Для обычных рейсов отображается маршрут между аэропортами, построенный по ортодромии, а для вертолетного режима карта фокусируется на Тверской области и доступных вертолетных площадках. На карте отображаются точки вылета и прилета, линия маршрута и маркер положения рейса по прогрессу выполнения.",
        "В списке рейсов реализованы фильтры по уровню риска, статусу решения, аэропорту и времени вылета. Отдельно предусмотрены аналитические блоки: количество рейсов, распределение по рискам, число запросов к метеорологу и другие показатели, помогающие диспетчеру быстро оценить общую загрузку и опасные направления.",
        "Для выбранного рейса реализована печатная форма отчета. В нее включаются тип воздушного судна, маршрут, время вылета и прилета, риск по этапам, факторы риска, статус решения диспетчера и пояснение. Это позволяет использовать систему не только как оперативный экран, но и как источник документирования принятого решения.",
        "Пользовательские процессы создания рейса и запроса данных у метеоролога дополнительно оформлены в виде BPMN-схем. Процесс создания рейса приведен на рисунке В.2, а процесс запроса метеорологу - на рисунке В.3. Схема перерасчета риска после автоматического или ручного обновления погодных данных приведена на рисунке В.4.",
    ],
    "2.8 Проектирование и реализация базы данных": [
        "Модель данных включает не только основную таблицу рейсов, но и таблицу истории flight_history. Основная сущность рейса хранит номер, тип воздушного судна, аэропорты вылета и прилета, время вылета и прибытия, частные и общий показатели риска, факторы риска, итоговую реализуемость и решение диспетчера. История хранит изменения риска и решений во времени.",
        "Поле aircraftType используется для различения самолетных и вертолетных рейсов. Это поле применяется при валидации номера, при оценке времени прибытия и при отображении рейса в интерфейсе. Такой подход расширяет предметную область проекта без усложнения основной структуры API.",
    ],
}


def add_test_cases(doc):
    add_plain(
        doc,
        "Для проверки работоспособности разработанной системы были подготовлены тест-кейсы, охватывающие основные пользовательские и отказоустойчивые сценарии.",
    )
    cases = [
        {
            "caption": "Таблица 1 - Проверка создания рейса с корректными данными",
            "priority": "Высокий",
            "name": "Создание рейса с корректными данными",
            "summary": "Проверка создания карточки рейса, расчета риска и отображения маршрута на карте",
            "steps": [
                "Выбрать тип воздушного судна «Самолет».",
                "Указать номер рейса SU1492.",
                "Выбрать аэропорт вылета SVO и аэропорт прибытия KGD.",
                "Указать плановое время вылета и нажать кнопку создания рейса.",
            ],
            "data": "Тип ВС: самолет; номер рейса: SU1492; маршрут: SVO - KGD; время вылета: будущая дата",
            "expected": "Рейс создан, маршрут отображен на карте, рассчитаны риск вылета, риск прибытия, риск маршрута и общий риск",
            "actual": "Рейс создан, риск рассчитан, факторы риска отображены",
            "post": "В списке рейсов появилась новая карточка рейса",
            "status": "Pass",
        },
        {
            "caption": "Таблица 2 - Проверка валидации некорректного номера рейса",
            "priority": "Высокий",
            "name": "Валидация некорректного номера рейса",
            "summary": "Проверка запрета создания рейса с номером, не соответствующим формату",
            "steps": [
                "Открыть форму создания рейса.",
                "Выбрать тип воздушного судна «Самолет».",
                "Ввести номер рейса 123ABC.",
                "Заполнить аэропорты и время вылета.",
                "Отправить форму создания рейса.",
            ],
            "data": "Тип ВС: самолет; номер рейса: 123ABC; маршрут: SVO - LED",
            "expected": "Система не создает рейс и выводит сообщение о корректном формате номера рейса",
            "actual": "Рейс не создан, отображено сообщение об ошибке формата номера",
            "post": "-",
            "status": "Pass",
        },
        {
            "caption": "Таблица 3 - Проверка запроса уточнения данных у метеоролога",
            "priority": "Средний",
            "name": "Запрос уточнения данных у метеоролога",
            "summary": "Проверка сценария взаимодействия диспетчера и метеоролога при необходимости уточнения погодных данных",
            "steps": [
                "Открыть карточку рейса в режиме диспетчера.",
                "Нажать кнопку запроса данных у метеоролога.",
                "Перейти в режим метеоролога.",
                "Заполнить METAR, TAF, данные о грозе, обледенении, ветре и видимости.",
                "Отправить ответ диспетчеру.",
            ],
            "data": "Рейс: SU1492; METAR: UUEE 251200Z 22008MPS 9999 SCT020 06/M01 Q1018; TAF: UUEE 251100Z 2512/2612 21007MPS 9999 BKN020",
            "expected": "Ответ метеоролога сохраняется, диспетчер получает уведомление, данные доступны для повторной оценки риска",
            "actual": "Ответ сохранен, уведомление отображено, сценарий завершен успешно",
            "post": "Запрос помечен как обработанный",
            "status": "Pass",
        },
        {
            "caption": "Таблица 4 - Проверка моделирования задержки рейса",
            "priority": "Высокий",
            "name": "Моделирование задержки рейса",
            "summary": "Проверка what-if расчета риска при переносе времени вылета",
            "steps": [
                "Открыть карточку рейса с рассчитанным риском.",
                "Открыть форму решения диспетчера.",
                "Выбрать решение DELAY.",
                "Указать задержку 60 минут.",
                "Нажать кнопку расчета what-if.",
            ],
            "data": "Рейс: SU1492; задержка: 60 минут; допустимый диапазон задержки: 5-360 минут",
            "expected": "Система показывает текущий риск, риск после задержки, изменение риска, новое время вылета и новое время прибытия",
            "actual": "What-if расчет выполнен, новые значения риска и времени отображены",
            "post": "Решение может быть сохранено диспетчером с указанием причины",
            "status": "Pass",
        },
    ]

    for case in cases:
        add_single_test_case_table(doc, case)


def add_single_test_case_table(doc, case):
    cap = doc.add_paragraph(style="Diplom Caption")
    run = cap.add_run(case["caption"])
    set_run_font(run)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    widths = (7.0, 10.0)

    def row(label, value, multiline=False):
        cells = table.add_row().cells
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        set_cell_text(cells[0], label, bold=True)
        if multiline:
            cells[1].text = ""
            for index, item in enumerate(value, start=1):
                append_cell_paragraph(cells[1], f"{index}. {item}")
        else:
            set_cell_text(cells[1], value)

    row("Приоритет тестирования", case["priority"])
    row("Наименование тестирования", case["name"])
    row("Резюме испытания", case["summary"])
    row("Шаги тестирования", case["steps"], multiline=True)
    row("Данные тестирования", case["data"])
    row("Ожидаемый результат", case["expected"])
    row("Фактический результат", case["actual"])
    row("Постусловия", case["post"])
    row("Статус", case["status"])


def classify(text):
    if text == "ВВЕДЕНИЕ":
        return ("heading", "Введение", 1, True)
    if text == "Заключение":
        return ("heading", text, 1, True)
    if re.fullmatch(r"\d+ [А-ЯA-Z].*", text):
        return ("heading", text, 1, True)
    if re.fullmatch(r"\d+\.\d+ .+", text):
        return ("heading", text, 2, False)
    if re.fullmatch(r"\d+\.\d+\.\d+ .+", text):
        return ("heading", text, 3, False)
    return ("body", text, None, False)


def add_body(doc):
    for raw in extract_clean_body():
        kind, text, level, page_break = classify(raw)
        if kind == "heading":
            add_heading(doc, text, level=level, page_break=page_break)
            for extra in EXTRA_BY_HEADING.get(text, []):
                if isinstance(extra, dict) and extra.get("type") == "formula":
                    add_formula(doc, extra["expression"], extra["number"], extra["explanations"])
                else:
                    add_plain(doc, extra)
            if text == "2.10 Тестирование и верификация корректности работы системы":
                add_test_cases(doc)
            continue

        if ": –" in text:
            intro, rest = text.split(": –", 1)
            add_plain(doc, intro + ":")
            for chunk in re.split(r";\s+–\s*", rest):
                cleaned = chunk.strip().rstrip(";")
                if cleaned:
                    add_bullet(doc, cleaned)
            continue

        if "; –" in text:
            for chunk in re.split(r";\s+–\s*", text):
                cleaned = chunk.strip().rstrip(";")
                if cleaned:
                    add_bullet(doc, cleaned)
            continue

        chunks = [chunk.strip() for chunk in re.split(r"\u2028", text) if chunk.strip()]
        if len(chunks) > 1:
            first = chunks[0]
            if not first.startswith(("–", "-", "•")):
                add_plain(doc, first)
                chunks = chunks[1:]
            for chunk in chunks:
                cleaned = re.sub(r"^[–\-•]\s*", "", chunk).strip()
                add_bullet(doc, cleaned)
            continue

        if text.startswith(("– ", "- ", "• ")):
            add_bullet(doc, re.sub(r"^[–\-•]\s*", "", text))
        else:
            add_plain(doc, text)


def add_literature(doc):
    add_heading(doc, "Список литературы и источников", level=1, page_break=True)
    sources = [
        "ГОСТ 19.101-77. Единая система программной документации. Виды программ и программных документов. - М.: Стандартинформ, 1978.",
        "ГОСТ 19.201-78. Единая система программной документации. Техническое задание на программу. Требования к содержанию и оформлению. - М.: Стандартинформ, 1979.",
        "ГОСТ 19.404-79. Единая система программной документации. Описание программы. Требования к содержанию и оформлению. - М.: Стандартинформ, 1980.",
        "ГОСТ 34.201-89. Информационная технология. Виды, комплектность и обозначения документов при создании автоматизированных систем. - М.: Изд-во стандартов, 1990.",
        "ГОСТ 34.602-89. Информационная технология. Техническое задание на создание автоматизированной системы. - М.: Изд-во стандартов, 1991.",
        "Блинов И. В., Соловьев А. А. Основы проектирования информационных систем. - М.: Академия, 2020. - 320 с.",
        "Воронин Е. Н. Введение в проектирование информационных систем. - М.: ФИЗМАТЛИТ, 2021. - 256 с.",
        "Горшков И. А. Проектирование баз данных: учебное пособие. - М.: Юрайт, 2022. - 314 с.",
        "Ермаков С. Ю. Информационные технологии в управлении бизнес-процессами. - М.: Инфра-М, 2021. - 368 с.",
        "Соловьев А. А., Терехов А. Н. Информационные системы и технологии. - М.: Горячая линия - Телеком, 2019. - 480 с.",
        "Spring. Spring Boot Reference Documentation [Электронный ресурс]. - URL: https://docs.spring.io/spring-boot/ - Дата обращения: 05.05.2026.",
        "Oracle. Java Platform, Standard Edition Documentation [Электронный ресурс]. - URL: https://docs.oracle.com/en/java/ - Дата обращения: 05.05.2026.",
        "React. React Documentation [Электронный ресурс]. - URL: https://react.dev/ - Дата обращения: 05.05.2026.",
        "PostgreSQL Global Development Group. PostgreSQL Documentation [Электронный ресурс]. - URL: https://www.postgresql.org/docs/ - Дата обращения: 05.05.2026.",
        "Docker Inc. Docker Documentation [Электронный ресурс]. - URL: https://docs.docker.com/ - Дата обращения: 05.05.2026.",
        "OpenWeatherMap. Weather API [Электронный ресурс]. - URL: https://openweathermap.org/api - Дата обращения: 05.05.2026.",
        "MDN Web Docs. Web Storage API [Электронный ресурс]. - URL: https://developer.mozilla.org/en-US/docs/Web/API/Web_Storage_API - Дата обращения: 05.05.2026.",
        "Hibernate. Hibernate ORM Documentation [Электронный ресурс]. - URL: https://hibernate.org/orm/documentation/ - Дата обращения: 05.05.2026.",
        "Vite. Vite Documentation [Электронный ресурс]. - URL: https://vite.dev/guide/ - Дата обращения: 05.05.2026.",
        "Leaflet. Leaflet Documentation [Электронный ресурс]. - URL: https://leafletjs.com/reference.html - Дата обращения: 05.05.2026.",
        "OpenStreetMap Foundation. OpenStreetMap Wiki [Электронный ресурс]. - URL: https://wiki.openstreetmap.org/ - Дата обращения: 05.05.2026.",
        "Mantine. Mantine Documentation [Электронный ресурс]. - URL: https://mantine.dev/ - Дата обращения: 05.05.2026.",
        "Spring Data. Spring Data JPA Reference Documentation [Электронный ресурс]. - URL: https://docs.spring.io/spring-data/jpa/reference/ - Дата обращения: 05.05.2026.",
    ]
    for i, source in enumerate(sources, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"{i}. {source}")
        set_run_font(run)


def add_appendix(doc):
    add_heading(doc, "Приложение А", level=1, page_break=True)
    add_centered(doc, "(справочное)")
    add_centered(doc, "Структура проекта", 12, True)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    widths = (5.2, 10.8)
    headers = ("Компонент", "Назначение")
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
        set_cell_width(table.rows[0].cells[idx], widths[idx])
    rows = [
        ("backend", "Серверная часть на Java и Spring Boot: REST API, бизнес-логика, расчет риска, работа с PostgreSQL."),
        ("frontend", "Клиентская часть на React: интерфейс диспетчера и метеоролога, карта маршрута, карточки рейсов."),
        ("docker-compose.yml", "Контейнерное развертывание backend, frontend и PostgreSQL."),
        ("airports-rf.json", "Справочник аэропортов и площадок, используемый для выбора маршрута и координат погодного запроса."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)
            set_cell_width(cells[idx], widths[idx])
    cap = doc.add_paragraph(style="Diplom Caption")
    cap.add_run("Таблица А.1 - Структура программного проекта")

    add_heading(doc, "Приложение Б", level=1, page_break=True)
    add_centered(doc, "(справочное)")
    add_centered(doc, "Инструкция по запуску проекта", 12, True)
    for item in [
        "Проверить наличие Docker и Docker Compose на рабочем компьютере.",
        "В корневом каталоге проекта выполнить команду docker compose up --build.",
        "Открыть пользовательский интерфейс frontend в браузере по адресу, указанному в настройках docker-compose.yml.",
        "Проверить доступность backend API и подключение к базе данных PostgreSQL.",
        "Создать тестовый рейс, выбрать аэропорты вылета и прибытия, выполнить обновление риска и проверить отображение рекомендации.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Приложение В", level=1, page_break=True)
    add_centered(doc, "(справочное)")
    add_centered(doc, "Графические схемы проекта", 12, True)
    add_plain(
        doc,
        "В приложении представлены схемы, разработанные в ходе проектирования веб-системы мониторинга и анализа погодных условий для авиационных рейсов. Схемы дополняют описание архитектуры, пользовательских сценариев и процесса перерасчета риска.",
    )
    for index, (path, caption, title, width, height) in enumerate(SCHEMES):
        if index > 0:
            doc.add_page_break()
        add_centered(doc, title, 12, True)
        if path.exists():
            add_picture_centered(doc, path, width=width, height=height)
        else:
            add_plain(doc, f"Файл схемы не найден: {path}", first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        cap = doc.add_paragraph(style="Diplom Caption")
        run = cap.add_run(caption)
        set_run_font(run)


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_assignment(doc)
    add_contents(doc)
    add_body(doc)
    add_literature(doc)
    add_appendix(doc)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
