from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/kotewa/my/diplom на проверку.docx")
USE_CASE_IMAGE = Path("/Users/kotewa/my/Диплом схемы/use-case.png")
NEW_CAPTION = "Рисунок 6 - Актуализированная диаграмма вариантов использования системы"


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.5)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)


def format_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def add_run_with_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run)


def ensure_detailed_use_case_text(doc: Document):
    marker_prefix = "Также была создана диаграмма вариантов использования."
    detail_first = (
        "Подробная структура пользовательских сценариев представлена на актуализированной диаграмме "
        "вариантов использования (Приложение, рисунок 6)."
    )

    if any(detail_first in p.text for p in doc.paragraphs):
        return

    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(marker_prefix):
            idx = i
            break
    if idx is None:
        raise RuntimeError("Не найден абзац для вставки описания use case")

    # Обновляем ссылку в исходном абзаце на новый рисунок.
    updated = doc.paragraphs[idx].text.replace("Рисунок 3", "рисунок 6")
    add_run_with_text(doc.paragraphs[idx], updated)
    format_body(doc.paragraphs[idx])

    details = [
        detail_first,
        (
            "Центральным вариантом использования является «Создание и планирование рейса». "
            "В рамках этого сценария диспетчер формирует карточку рейса, выбирает маршрут, "
            "задает плановое время и инициирует первичную оценку метеоусловий."
        ),
        (
            "Диаграмма показывает связь include между вариантом «Создание и планирование рейса» и "
            "вариантом «Принятие решения о времени осуществления рейса». Это означает, что после "
            "планирования система обязательно поддерживает этап принятия решения: разрешить рейс, "
            "задержать его или отменить при повышенном уровне риска."
        ),
        (
            "Вторая связь include связывает планирование с вариантом «Просмотр сводной аналитики рейсов». "
            "В аналитическом блоке диспетчер видит обобщенные показатели: общее количество рейсов, "
            "распределение по уровням риска, статистику решений и число запросов метеорологу. "
            "Это позволяет оценивать не только один рейс, но и общую оперативную картину смены."
        ),
        (
            "Вариант «Выдача метеорологических данных диспетчеру по запросу» выполняется актором "
            "«Метеоролог». По запросу диспетчера метеоролог передает уточняющие параметры "
            "(METAR, TAF, грозовая обстановка, обледенение, ветер, видимость и условия по маршруту), "
            "после чего система повторно оценивает риск и обновляет рекомендации."
        ),
        (
            "Отдельный вариант «Создание отчетной документации» закрепляет этап документирования. "
            "После анализа и принятия решения диспетчер формирует отчет по рейсу, в который входят "
            "основные данные рейса, оценка рисков, ключевые метеофакторы, итоговое решение и история событий."
        ),
        (
            "Таким образом, диаграмма на рисунке 6 отражает полный цикл работы системы: "
            "от планирования рейса и получения метеоданных до аналитики, принятия решения и формирования отчетности."
        ),
    ]

    insert_at = idx + 1
    for text in details:
        p = doc.paragraphs[insert_at].insert_paragraph_before(text)
        format_body(p)
        insert_at += 1


def ensure_appendix_image(doc: Document):
    if any(p.text.strip() == NEW_CAPTION for p in doc.paragraphs):
        return

    doc.add_page_break()
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.first_line_indent = Cm(0)
    pic.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pic.paragraph_format.space_before = Pt(0)
    pic.paragraph_format.space_after = Pt(0)
    run = pic.add_run()
    run.add_picture(str(USE_CASE_IMAGE), width=Cm(16.0))

    cap = doc.add_paragraph(NEW_CAPTION)
    format_caption(cap)


def main():
    doc = Document(DOCX_PATH)
    ensure_detailed_use_case_text(doc)
    ensure_appendix_image(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
