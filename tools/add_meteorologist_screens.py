from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/kotewa/my/diplom на проверку.docx")
SCREENS_DIR = Path("/Users/kotewa/my/скрины")

ITEMS = [
    {
        "mention": "Рабочее место метеоролога с перечнем входящих запросов от диспетчера показано на рисунке 15. Такой экран позволяет выбрать конкретный рейс и перейти к подготовке уточненных метеорологических данных.",
        "image": SCREENS_DIR / "Снимок экрана 2026-05-13 в 21.29.58.png",
        "caption": "Рисунок 15 - Рабочее место метеоролога со списком запросов",
        "width": 16.0,
    },
    {
        "mention": "Форма ввода метеорологических данных для ответа диспетчеру представлена на рисунке 16. В ней метеоролог заполняет METAR, TAF, сведения о грозовой обстановке, риске обледенения, ветре, видимости и условиях на маршруте.",
        "image": SCREENS_DIR / "Снимок экрана 2026-05-13 в 21.30.13.png",
        "caption": "Рисунок 16 - Форма ввода данных метеорологом",
        "width": 16.0,
    },
]


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_after=0):
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.5 if first_indent else 0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        set_run_font(run)


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    new_p.addnext(result._p)
    paragraph._p.getparent().remove(new_p)
    if text:
        run = result.add_run(text)
        set_run_font(run)
    return result


def insert_image_block(after_paragraph, mention_text, image_path, caption_text, width_cm):
    mention = paragraph_after(after_paragraph, mention_text)
    format_paragraph(mention)

    pic = paragraph_after(mention)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.first_line_indent = Cm(0)
    pic.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pic.paragraph_format.space_before = Pt(6)
    pic.paragraph_format.space_after = Pt(0)
    pic.add_run().add_picture(str(image_path), width=Cm(width_cm))

    caption = paragraph_after(pic, caption_text)
    format_paragraph(caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=6)
    return caption


def contains(doc, text):
    return any(paragraph.text.strip() == text for paragraph in doc.paragraphs)


def main():
    doc = Document(DOCX_PATH)
    if all(contains(doc, item["caption"]) for item in ITEMS):
        return

    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Кнопка «Запросить данные у метеоролога» отображается"):
            anchor = paragraph
            break
    if anchor is None:
        raise RuntimeError("Meteorologist scenario paragraph was not found")

    current = anchor
    for item in ITEMS:
        if contains(doc, item["caption"]):
            continue
        current = insert_image_block(
            current,
            item["mention"],
            item["image"],
            item["caption"],
            item["width"],
        )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
