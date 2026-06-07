from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("/Users/kotewa/my/diplom на проверку.docx")

REQUIRED_GOSTS = [
    "ГОСТ Р 57240-2016. Воздушный транспорт. Менеджмент безопасности авиационной деятельности в гражданской авиации. Основные положения [Электронный ресурс]. - URL: https://protect.gost.ru/gost/details/6094796d-8ed9-48b0-b222-09e48ffc05f0 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 57908-2017. Воздушный транспорт. Система менеджмента безопасности авиационной деятельности. База данных. Авиационные риски по реализации системы оценки безопасности полетов при обеспечении воздушного движения [Электронный ресурс]. - URL: https://internet-law.ru/gosts/gost/65810/ - Дата обращения: 16.05.2026.",
    "ГОСТ Р 58712-2019. Автоматизированная метеорологическая измерительная система. Общие технические требования [Электронный ресурс]. - URL: https://www.standards.ru/print.aspx?control=27&id=8096369&print=yes - Дата обращения: 16.05.2026.",
    "ГОСТ 22.1.01-2023. Безопасность в чрезвычайных ситуациях. Мониторинг и прогнозирование. Основные положения [Электронный ресурс]. - URL: https://www.gostinfo.ru/catalog/Details/?id=7476806 - Дата обращения: 16.05.2026.",
    "ГОСТ Р 22.1.07-2023. Безопасность в чрезвычайных ситуациях. Мониторинг и прогнозирование опасных метеорологических явлений и процессов. Общие требования [Электронный ресурс]. - URL: https://base.garant.ru/408555345/ - Дата обращения: 16.05.2026.",
]


def set_run_font(run, size=14):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def format_biblio_paragraph(paragraph):
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    # Force plain paragraph (no automatic numbering from Word list styles)
    ppr = paragraph._element.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)
    for run in paragraph.runs:
        set_run_font(run)


def clean_entry(text: str) -> str:
    text = text.replace("\u2022", " ").strip()
    text = re.sub(r"^\d+\.\s*", "", text).strip()
    text = re.sub(r"\s+", " ", text).strip()
    return text


def paragraph_after(paragraph, text=""):
    from docx.oxml import OxmlElement

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    new_p.addnext(result._p)
    paragraph._p.getparent().remove(new_p)
    if text:
        result.add_run(text)
    return result


def main():
    doc = Document(DOCX_PATH)

    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().lower()
        if start is None and t.startswith("список литературы"):
            start = i
        elif start is not None and t.startswith("приложение"):
            end = i
            break

    if start is None:
        raise RuntimeError("Не найден раздел списка литературы")
    if end is None:
        end = len(doc.paragraphs)

    raw_entries = []
    for p in doc.paragraphs[start + 1 : end]:
        t = clean_entry(p.text)
        if not t:
            continue
        raw_entries.append(t)

    # Remove duplicates while preserving order.
    entries = []
    seen = set()
    for item in raw_entries:
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        entries.append(item)

    # Ensure required aviation/weather/risk GOSTs are present.
    for gost in REQUIRED_GOSTS:
        key = gost.lower()
        if key not in seen:
            entries.append(gost)
            seen.add(key)

    # Remove old bibliography paragraphs.
    for p in list(doc.paragraphs[start + 1 : end]):
        p._element.getparent().remove(p._element)

    anchor = doc.paragraphs[start]
    for idx, item in enumerate(entries, start=1):
        new_p = paragraph_after(anchor, f"{idx}. {item}")
        format_biblio_paragraph(new_p)
        anchor = new_p

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
